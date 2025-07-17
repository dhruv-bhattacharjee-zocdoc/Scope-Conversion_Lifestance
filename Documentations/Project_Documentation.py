from docx import Document
from docx.shared import Pt, Inches, RGBColor
from pygments import highlight
from pygments.lexers import PythonLexer
from pygments.formatters import ImageFormatter
from io import BytesIO
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pygments.token import Token
from pygments import lex
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Data for documentation
workflow = [
    "Data Extraction: The main script (_main_1.py) reads the input Excel file and uses various helper scripts to extract fields like names, NPIs, specialties, etc.",
    "Template Population: Extracted data is mapped to a template Excel file, creating a new output file with the required structure.",
    "Dropdowns and Formulas: The script adds dropdowns and formulas for data validation and reference, using both built-in logic and helper scripts.",
    "Sheet Management: Additional sheets (like ValidationAndReference and Location) are copied or generated, and further dropdowns/formulas are added.",
    "Finalization: The output file is saved and opened for the user, ready for review or further editing."
]

files = [
    {
        "name": "_main_1.py",
        "purpose": "This is the central orchestrator of the workflow. It coordinates the entire process, from reading the input Excel file to producing the final output. The script calls various helper modules to extract, clean, and transform data, then maps this data into a template structure. It also manages the addition of dropdowns, formulas, and the integration of auxiliary sheets, ensuring the output is ready for downstream use and validation.",
        "snippet": "# Extract name and gender data using Name.py\nextracted_rows = extract_name_gender(input_file)\n# Extract NPI data using Npi.py\nnpi_list = extract_npi(input_file)\n# ... (other extractors)\n# Call Location.py to generate the Location sheet\nsubprocess.run(['python', 'Location.py'], check=True)\n# Add dropdowns and formulas\napply_provider_dropdowns(output_file, dropdown_specs)\napply_provider_formulas(output_file, formula_specs)"
    },
    {
        "name": "Location.py",
        "purpose": "This script is responsible for generating the Location sheet in the output Excel file. It copies relevant sheets from the template, standardizes and cleans address data, and ensures that each location is properly categorized (e.g., splitting 'Both' into 'Virtual' and 'In Person'). The script also manages the transfer of zip codes and other location-specific fields, ensuring consistency and completeness in the final output.",
        "snippet": "# Duplicate rows for 'Both' location types\nfor i, row in reversed(rows_to_duplicate):\n    ws_location.delete_rows(i)\n    # Insert two new rows: one with 'Virtual', one with 'In Person'\n    new_row_virtual = list(row)\n    new_row_virtual[location_type_idx] = 'Virtual'\n    new_row_inperson = list(row)\n    new_row_inperson[location_type_idx] = 'In Person'\n    ws_location.insert_rows(i)\n    for col_idx, value in enumerate(new_row_inperson, start=1):\n        ws_location.cell(row=i, column=col_idx, value=value)\n    ws_location.insert_rows(i)\n    for col_idx, value in enumerate(new_row_virtual, start=1):\n        ws_location.cell(row=i, column=col_idx, value=value)"
    },
    {
        "name": "Name.py",
        "purpose": "Handles the extraction of provider names and gender from the input Excel file. It ensures that the gender field is standardized (e.g., mapping 'Prefer not to say' to 'Not Applicable') and that all relevant name fields are captured accurately. This module is crucial for maintaining data integrity and consistency in the provider records.",
        "snippet": "def extract_name_gender(input_file):\n    # ...\n    for row in ws_in.iter_rows(min_row=2, values_only=True):\n        extracted = {col: row[input_indices[col]] for col in ['First Name', 'Last Name', 'Gender']}\n        if extracted['Gender'] == 'Prefer not to say':\n            extracted['Gender'] = 'Not Applicable'\n        extracted_rows.append(extracted)\n    return extracted_rows"
    },
    {
        "name": "Npi.py",
        "purpose": "Extracts National Provider Identifier (NPI) numbers from the input file. This script ensures that each provider’s unique NPI is accurately retrieved and mapped, which is essential for provider identification and compliance with healthcare data standards.",
        "snippet": "def extract_npi(input_file):\n    # ...\n    for row in ws_in.iter_rows(min_row=2, values_only=True):\n        npi_list.append(row[npi_idx])\n    return npi_list"
    },
    {
        "name": "Headshot.py",
        "purpose": "Responsible for extracting headshot URLs from the input data. This allows the output file to include direct links to provider photos, which can be used for display in downstream systems or directories.",
        "snippet": "def extract_headshot(input_file):\n    # ...\n    for row in ws_in.iter_rows(min_row=2, values_only=True):\n        headshot_list.append(row[headshot_idx])\n    return headshot_list"
    },
    {
        "name": "professional_suffix.py",
        "purpose": "Extracts professional suffixes (such as MD, PhD, etc.) from the input and applies them to the output file. It also manages the addition of dropdowns for these suffixes, ensuring users can select from standardized options when editing the output.",
        "snippet": "def extract_professional_suffix(input_file):\n    # ... (logic to extract suffixes)\n# Adds dropdowns for professional suffixes\nadd_professional_suffix_dropdowns(output_file)"
    },
    {
        "name": "Specialty.py",
        "purpose": "Extracts provider specialty information from the input file and manages the application of specialty dropdowns in the output. This ensures that specialty data is both accurate and validated against a reference list, supporting downstream reporting and analytics.",
        "snippet": "def extract_specialty(input_file):\n    # ...\n    for row in ws_in.iter_rows(min_row=2, values_only=True):\n        specialty_list.append(row[specialty_idx])\n    return specialty_list"
    },
    {
        "name": "PatientsAccepted.py",
        "purpose": "Extracts information about the types of patients accepted by each provider (e.g., Adult, Pediatric, Both) and sets up the corresponding dropdowns in the output file. This helps ensure that patient acceptance data is standardized and easy to update.",
        "snippet": "def extract_patients_accepted(input_file):\n    # ... (extract logic)\nset_patients_accepted_dropdown(output_file)"
    },
    {
        "name": "Education.py",
        "purpose": "Handles the extraction of education and school information for each provider. This module ensures that educational backgrounds are accurately captured and mapped, supporting credentialing and provider profiling.",
        "snippet": "def extract_education(input_file):\n    # ... (extract logic)\n    return education_list"
    },
    {
        "name": "Professional_statement.py",
        "purpose": "Extracts provider bios or professional statements from the input file. This information is important for provider directories and public profiles, giving context about each provider’s background and philosophy.",
        "snippet": "def extract_professional_statement(input_file):\n    # ...\n    for row in ws_in.iter_rows(min_row=2, values_only=True):\n        bio_list.append(row[bio_idx])\n    return bio_list"
    },
    {
        "name": "Board_certification.py",
        "purpose": "Extracts board certification and subspecialty information for each provider. It also manages the addition of dropdowns for board certifications, ensuring that only valid certifications are selectable in the output.",
        "snippet": "def extract_board_certification(input_file):\n    # ... (extract logic)\nset_board_certification_dropdown(output_file)"
    },
    {
        "name": "optoutrating.py",
        "purpose": "Adds a dropdown for the 'Opt Out of Ratings' field in the output file. This allows providers to indicate whether they wish to be excluded from ratings, supporting privacy and compliance requirements.",
        "snippet": "def set_opt_out_of_ratings_dropdown(output_file):\n    # ... (dropdown logic)"
    },
    {
        "name": "ESF.py",
        "purpose": "Adds a dropdown for the 'Enterprise Scheduling Flag' in the output file. This field is used to indicate whether a provider participates in enterprise-level scheduling, which can affect appointment availability and system integration.",
        "snippet": "def set_enterprise_scheduling_flag_dropdown(output_file):\n    # ... (dropdown logic)"
    },
    {
        "name": "Langauge.py",
        "purpose": "Extracts the languages spoken by each provider from the input file and sets up language dropdowns in the output. This ensures that language data is standardized and can be used for filtering or matching providers to patient needs.",
        "snippet": "def extract_languages(input_file):\n    # ...\n    for row in ws_in.iter_rows(min_row=2, values_only=True):\n        # ... (split and assign languages)\n    return lang1_list, lang2_list"
    },
    {
        "name": "provider_dropdowns.py",
        "purpose": "Applies a wide range of dropdowns and formulas to the Provider sheet in the output file, based on a list of specifications. This script centralizes the logic for data validation and formula application, making the output robust and user-friendly.",
        "snippet": "apply_provider_dropdowns(output_file, dropdown_specs)\napply_provider_formulas(output_file, formula_specs)"
    },
    {
        "name": "specialtydropdown.py",
        "purpose": "Adds specialty dropdowns to the output file using validation references. This ensures that specialty selections are always consistent with the reference data, reducing errors and improving data quality.",
        "snippet": "add_specialty_valref_dropdowns(output_file)"
    },
    {
        "name": "_status _check.py",
        "purpose": "Checks the status and completeness of the output file by validating columns and grouping related fields. This script helps ensure that the final output meets all structural and data requirements before delivery or further processing.",
        "snippet": "# Group columns by base name (e.g., 'Specialty 1', 'Specialty 2' -> 'Specialty')\nfor col_idx, col_name in enumerate(header):\n    base = re.sub(r'\\s*\\d+$', '', str(col_name))\n    if base not in col_groups:\n        col_groups[base] = []\n    col_groups[base].append((col_name, col_idx))"
    }
]

def code_to_image(code_text):
    formatter = ImageFormatter(
        font_name='Consolas',
        line_numbers=False,
        style='monokai',  # Modern editor-like look
        image_format='PNG',
        line_pad=2,
        font_size=14      # Uniform font size
    )
    img_data = highlight(code_text, PythonLexer(), formatter)
    return BytesIO(img_data)

def add_code_block(doc, code_text, language_label=None):
    # Add language label if provided
    if language_label:
        label_para = doc.add_paragraph(language_label)
        label_para.style = doc.styles['Normal']
        label_para.runs[0].font.size = Pt(8)
        label_para.runs[0].font.italic = True
    # Add code block with darker grey background
    para = doc.add_paragraph()
    # Set background shading for the paragraph
    ppr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'D9D9D9')  # Darker grey
    ppr.append(shd)
    # Tokenize code and add runs with color
    for token, value in lex(code_text, PythonLexer()):
        run = para.add_run(value)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        # Set color based on token type
        if token in Token.Keyword:
            run.font.color.rgb = RGBColor(0, 0, 205)  # Blue
        elif token in Token.Comment:
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        elif token in Token.String:
            run.font.color.rgb = RGBColor(163, 21, 21)  # Red
        elif token in Token.Number:
            run.font.color.rgb = RGBColor(128, 0, 128)  # Purple
        else:
            run.font.color.rgb = RGBColor(51, 51, 51)  # Default dark grey
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(6)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

doc = Document()

# Add left-aligned title with a blue horizontal line below
title_para = doc.add_paragraph('Lifestance Scope Conversion Code Documentation')
title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
title_run = title_para.runs[0]
title_run.font.size = Pt(18)
title_run.font.bold = False
title_run.font.name = 'Calibri'
title_run.font.color.rgb = RGBColor(0x17, 0x36, 0x5d)  # #17365d
# Add a horizontal line (bottom border)
p = title_para._p
pPr = p.get_or_add_pPr()
borders = pPr.find(qn('w:pBdr'))
if borders is None:
    borders = OxmlElement('w:pBdr')
    pPr.append(borders)
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '12')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '2F5496')  # Blue color
borders.append(bottom)

# Helper to set font for paragraphs
from docx.shared import RGBColor

def set_paragraph_font(para, size_pt, color_rgb=(0,0,0)):
    for run in para.runs:
        run.font.size = Pt(size_pt)
        run.font.name = 'Calibri'
        run.font.color.rgb = RGBColor(*color_rgb)

# Project Overview
# ----------------
doc.add_heading('Project Overview', level=1)
set_paragraph_font(doc.paragraphs[-1], 11, (0x17, 0x36, 0x5d))

doc.add_heading('Problem Statement', level=2)
doc.add_paragraph(
    "The project addresses the need to efficiently transpose and standardize provider data from various Excel sources into a unified, validated format. This is essential for downstream processes such as reporting, analytics, and integration with other healthcare systems. Manual data handling is error-prone and time-consuming, necessitating an automated, robust solution."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

doc.add_heading('Achieved Solution', level=2)
doc.add_paragraph(
    "A modular Python-based solution was developed to automate the extraction, transformation, and validation of provider data. The system leverages multiple helper scripts to extract specific fields, applies data validation and dropdowns, and generates a ready-to-use output file. The approach ensures data integrity, reduces manual effort, and supports scalability for future requirements."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

doc.add_heading('Impact', level=2)
doc.add_paragraph(
    "The automated workflow has significantly improved data quality and processing speed. It minimizes human errors, ensures compliance with data standards, and provides a flexible framework for future enhancements. The output is structured for easy review and integration, supporting both operational and analytical needs."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

# System Architecture & Workflow
# -----------------------------
doc.add_heading('System Architecture & Workflow', level=1)
set_paragraph_font(doc.paragraphs[-1], 11, (0x17, 0x36, 0x5d))

doc.add_heading('Workflow Steps', level=2)
for step in workflow:
    para = doc.add_paragraph(step, style='List Number')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    set_paragraph_font(para, 9)

doc.add_heading('Main Python Files', level=2)
for file in files:
    doc.add_heading(file['name'], level=3)
    set_paragraph_font(doc.paragraphs[-1], 11, (0x17, 0x36, 0x5d))
    para = doc.add_paragraph('Purpose:', style='Intense Quote')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    para = doc.add_paragraph(file['purpose'])
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    set_paragraph_font(para, 9)
    para = doc.add_paragraph('Key Code Snippet:', style='Intense Quote')
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    add_code_block(doc, file['snippet'], language_label='Python')
    explanation = f"This code snippet demonstrates how the script handles: {file['purpose']}"
    para = doc.add_paragraph()
    set_paragraph_font(para, 9)
    run = para.add_run(explanation)
    # No italics
    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

# Examples
# --------
doc.add_heading('Examples', level=1)
set_paragraph_font(doc.paragraphs[-1], 11, (0x17, 0x36, 0x5d))

doc.add_heading('Example: Data Extraction', level=2)
doc.add_paragraph(
    "The system extracts provider names, NPIs, specialties, and other fields from the input Excel file using dedicated helper scripts. Each script is responsible for a specific data domain, ensuring modularity and ease of maintenance."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

doc.add_heading('Example: Location Sheet Generation', level=2)
doc.add_paragraph(
    "The Location.py script processes and standardizes address data, handling special cases such as 'Both' location types by splitting them into 'Virtual' and 'In Person'. This ensures accurate categorization and downstream usability."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

doc.add_heading('Example: Specialty Dropdowns', level=2)
doc.add_paragraph(
    "Specialty.py and specialtydropdown.py manage the extraction and validation of provider specialties. Dropdowns are applied to ensure only valid specialties are selectable, supporting data quality and consistency."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

# Validation & Output
# ------------------
doc.add_heading('Validation & Output', level=1)
set_paragraph_font(doc.paragraphs[-1], 11, (0x17, 0x36, 0x5d))

doc.add_heading('Output Structure', level=2)
doc.add_paragraph(
    "The output Excel file is structured to match the required template, with all necessary fields, dropdowns, and validation rules applied. Additional sheets, such as ValidationAndReference and Location, are included as needed."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

doc.add_heading('Validation Checks', level=2)
doc.add_paragraph(
    "The _status _check.py script validates the completeness and correctness of the output file. It groups related columns, checks for missing or inconsistent data, and ensures the final deliverable meets all requirements."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

# Conclusion
# ----------
doc.add_heading('Conclusion', level=1)
set_paragraph_font(doc.paragraphs[-1], 11, (0x17, 0x36, 0x5d))
doc.add_paragraph(
    "This solution provides a scalable, maintainable, and robust approach to provider data transposition and validation. By automating key processes and enforcing data standards, it supports both current operational needs and future growth."
).alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
set_paragraph_font(doc.paragraphs[-1], 9)

output_path = r'C:/Users/dhruv.bhattacharjee/Desktop/PDO Data Transposition/Scope Conversion_Lifestance/Documentations/Code Documentation.docx'
doc.save(output_path)

# Open the generated documentation automatically (Windows only)
import os
os.startfile(output_path)
print('Word documentation generated as Code Documentation.docx') 